"""本地文档文本抽取。

该模块在调用大模型前把上传文件转换为可读文本，避免把 Word/PDF 的二进制
内容直接交给 LLM。外部转换工具可用时优先使用；不可用时使用标准库兜底。
"""

from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"
_WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_text(file_name: str, file_type: str, content: bytes, fallback_text: str = "") -> str:
    """按文件类型抽取可读文本；失败时返回前端提供的 fallback 文本。"""
    name = str(file_name or "").strip()
    mime = str(file_type or "").strip()
    data = content or b""

    if _is_docx(name, mime, data):
        text = extract_docx_text(data)
        if text:
            return text
    if _is_doc(name, mime, data):
        text = extract_doc_text(data)
        if text:
            return text
    return str(fallback_text or "").strip()


def _is_docx(file_name: str, file_type: str, content: bytes) -> bool:
    return (
        file_name.lower().endswith(".docx")
        or file_type == DOCX_MIME
        or _looks_like_docx(content)
    )


def _is_doc(file_name: str, file_type: str, content: bytes) -> bool:
    return (
        file_name.lower().endswith(".doc")
        or file_type == DOC_MIME
        or content.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
    )


def _looks_like_docx(content: bytes) -> bool:
    if not content.startswith(b"PK"):
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            return "word/document.xml" in archive.namelist()
    except zipfile.BadZipFile:
        return False


def extract_docx_text(content: bytes) -> str:
    if not content:
        return ""
    docx_text = _extract_docx_text_with_python_docx(content)
    if docx_text:
        return docx_text
    return _extract_docx_text_with_xml(content)


def _extract_docx_text_with_python_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except Exception:
        return ""
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                parts.append(text)
    return _clean_text("\n".join(parts))


def _extract_docx_text_with_xml(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            names = [
                name for name in archive.namelist()
                if name.startswith("word/")
                and name.endswith(".xml")
                and (
                    name == "word/document.xml"
                    or name.startswith("word/header")
                    or name.startswith("word/footer")
                    or name.startswith("word/footnotes")
                    or name.startswith("word/endnotes")
                )
            ]
            paragraphs: list[str] = []
            for name in sorted(names, key=lambda item: (item != "word/document.xml", item)):
                root = ElementTree.fromstring(archive.read(name))
                paragraphs.extend(_paragraph_text(root))
            return _clean_text("\n".join(paragraphs))
    except (zipfile.BadZipFile, KeyError, ElementTree.ParseError):
        return ""


def _paragraph_text(root: ElementTree.Element) -> list[str]:
    paragraphs: list[str] = []
    for paragraph in root.iter(f"{_WORD_NS}p"):
        parts: list[str] = []
        for node in paragraph.iter():
            if node.tag == f"{_WORD_NS}t" and node.text:
                parts.append(node.text)
            elif node.tag == f"{_WORD_NS}tab":
                parts.append("\t")
            elif node.tag == f"{_WORD_NS}br":
                parts.append("\n")
        text = "".join(parts).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def extract_doc_text(content: bytes) -> str:
    if not content:
        return ""
    converted = _extract_doc_with_external_tool(content)
    if converted:
        return converted
    return _extract_doc_with_heuristics(content)


def _extract_doc_with_external_tool(content: bytes) -> str:
    if shutil.which("antiword"):
        return _run_stdin_tool(["antiword", "-"], content)
    if shutil.which("catdoc"):
        return _run_stdin_tool(["catdoc", "-"], content)
    office = shutil.which("libreoffice") or shutil.which("soffice")
    if office:
        return _run_libreoffice_text_export(office, content)
    return ""


def _run_stdin_tool(command: list[str], content: bytes) -> str:
    try:
        completed = subprocess.run(
            command,
            input=content,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
            timeout=30,
        )
    except (OSError, subprocess.TimeoutExpired):
        return ""
    if completed.returncode != 0:
        return ""
    return _clean_text(completed.stdout.decode("utf-8", "replace"))


def _run_libreoffice_text_export(office: str, content: bytes) -> str:
    with tempfile.TemporaryDirectory(prefix="sim_doc_") as tmp:
        input_path = Path(tmp) / "input.doc"
        input_path.write_bytes(content)
        try:
            completed = subprocess.run(
                [
                    office,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    tmp,
                    str(input_path),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=False,
                timeout=60,
            )
        except (OSError, subprocess.TimeoutExpired):
            return ""
        if completed.returncode != 0:
            return ""
        output_path = Path(tmp) / "input.txt"
        if not output_path.exists():
            return ""
        return _clean_text(output_path.read_text(encoding="utf-8", errors="replace"))


def _extract_doc_with_heuristics(content: bytes) -> str:
    candidates = [
        content.decode("utf-16le", "ignore"),
        content.decode("utf-8", "ignore"),
        content.decode("gb18030", "ignore"),
    ]
    pieces: list[str] = []
    for candidate in candidates:
        pieces.extend(_readable_runs(candidate))
    unique = list(dict.fromkeys(pieces))
    cjk_pieces = [piece for piece in unique if _cjk_count(piece) >= 3]
    return _clean_text("\n".join(cjk_pieces or unique))


def _readable_runs(text: str) -> list[str]:
    normalized = text.replace("\x00", " ")
    pattern = re.compile(r"[\u4e00-\u9fffA-Za-z0-9，。；：、（）()《》<>/%+\-_\s]{4,}")
    return [match.group(0).strip() for match in pattern.finditer(normalized) if match.group(0).strip()]


def _cjk_count(text: str) -> int:
    return sum(1 for char in text if "\u4e00" <= char <= "\u9fff")


def _clean_text(text: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in str(text or "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
